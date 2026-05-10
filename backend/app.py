import os, re, json, io, shutil, tempfile
from pathlib import Path
from datetime import datetime
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import groq

load_dotenv()

app = Flask(__name__)
CORS(app)

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
FRONTEND_DIR = Path(__file__).parent.parent / "frontend"
UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)
MAX_DOCX_SIZE = 15 * 1024 * 1024  # 15MB


@app.route('/app')
@app.route('/app/')
def serve_app():
    resp = send_from_directory(str(FRONTEND_DIR), 'index.html')
    resp.headers['Cache-Control'] = 'no-store'
    return resp


# ── helpers ──────────────────────────────────────────────────────────────────

def extract_date_from_transcript(text: str) -> str:
    """Try to find a date in the transcript. Returns yyyymmdd or ''."""
    patterns = [
        r'\b(\d{4})[/-](\d{1,2})[/-](\d{1,2})\b',          # 2026-03-27
        r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{4})\b',          # 27/03/2026
        r'\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(\d{4})\b',
        r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b',
    ]
    months = {'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,
              'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12,
              'january':1,'february':2,'march':3,'april':4,'june':6,
              'july':7,'august':8,'september':9,'october':10,'november':11,'december':12}
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            g = m.groups()
            try:
                if re.match(r'\d{4}', g[0]):
                    y, mo, d = int(g[0]), int(g[1]), int(g[2])
                elif g[1].lower() in months:
                    d, mo, y = int(g[0]), months[g[1].lower()], int(g[2])
                else:
                    d, mo, y = int(g[0]), int(g[1]), int(g[2])
                return f"{y:04d}{mo:02d}{d:02d}"
            except Exception:
                continue
    return ""


def parse_meeting_info(text: str) -> dict:
    """Parse meeting name, date, time from pasted calendar text."""
    result = {}

    HARI_ID = ['Minggu','Senin','Selasa','Rabu','Kamis','Jumat','Sabtu']
    BULAN_ID = ['Januari','Februari','Maret','April','Mei','Juni',
                'Juli','Agustus','September','Oktober','November','Desember']
    MONTHS_EN = {'january':1,'february':2,'march':3,'april':4,'may':5,'june':6,
                 'july':7,'august':8,'september':9,'october':10,'november':11,'december':12,
                 'jan':1,'feb':2,'mar':3,'apr':4,'jun':6,'jul':7,'aug':8,
                 'sep':9,'oct':10,'nov':11,'dec':12}

    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]

    # Meeting name: strip prefix like [Book], [Booking], etc.
    if lines:
        name = re.sub(r'^\[.*?\]\s*', '', lines[0]).strip()
        result['meeting_name'] = name

    # Date + time line: "Thursday, May 7⋅3:30 – 4:30pm" or "Thursday, May 7 · 3:30 – 4:30pm"
    for line in lines[1:]:
        # Normalize separators
        line_n = line.replace('⋅','·').replace('•','·')
        m = re.match(
            r'(\w+),\s+(\w+)\s+(\d{1,2})(?:\s*·\s*(\d{1,2}:\d{2})(am|pm)?\s*[–-]\s*(\d{1,2}:\d{2})(am|pm)?)?',
            line_n, re.IGNORECASE
        )
        if m:
            day_en, month_en, day_num = m.group(1), m.group(2), int(m.group(3))
            mo = MONTHS_EN.get(month_en.lower())
            if mo:
                # MOM = past meeting; if month hasn't occurred yet this year, use previous year
                now = datetime.now()
                year = now.year if mo <= now.month else now.year - 1
                d = datetime(year, mo, day_num)
                hari_id = HARI_ID[d.weekday() + 1 if d.weekday() < 6 else 0]
                # weekday() Mon=0..Sun=6, we need Sun=0..Sat=6
                hari_id = HARI_ID[d.isoweekday() % 7]
                result['meeting_date'] = f"{hari_id}, {day_num} {BULAN_ID[mo-1]} {year}"
                result['meeting_date_iso'] = f"{year}-{mo:02d}-{day_num:02d}"

            # Time
            if m.group(4) and m.group(6):
                def to24(t, suffix):
                    h, mi = map(int, t.split(':'))
                    if suffix:  # has am/pm
                        is_pm = 'pm' in suffix
                        if is_pm and h != 12:
                            h += 12
                        elif not is_pm and h == 12:
                            h = 0
                    # else: no am/pm, use as-is (24h format)
                    return f"{h:02d}:{mi:02d}"
                start_suffix = (m.group(5) or '').lower()
                end_suffix = (m.group(7) or '').lower()
                # If only end has am/pm, apply to both
                if end_suffix and not start_suffix:
                    start_suffix = end_suffix
                result['meeting_time_start'] = to24(m.group(4), start_suffix)
                result['meeting_time_end'] = to24(m.group(6), end_suffix)
            break

    return result


@app.route('/api/parse-meeting-info', methods=['POST'])
def api_parse_meeting_info():
    body = request.get_json()
    text = body.get('text', '')
    return jsonify(parse_meeting_info(text))


def generate_mom_with_ai(transcript: str, meeting_name: str) -> dict:
    """Call Groq to generate MOM items from transcript."""
    if not GROQ_API_KEY:
        return {"error": "GROQ_API_KEY tidak ditemukan di backend/.env"}

    client = groq.Groq(api_key=GROQ_API_KEY)
    prompt = f"""You are a meeting secretary. Analyze this meeting transcript and extract Minutes of Meeting (MOM) data.

IMPORTANT LANGUAGE RULES for the "items" field:
- Detect the language of the transcript (English or Indonesian).
- If English: translate each issue/action item to Indonesian. Translate word by word and as a whole sentence naturally. If a term has no good Indonesian equivalent, keep the original term as-is.
- If Indonesian: keep as-is, but fix any typos or awkward phrasing.
- All issue text in the output MUST be in Indonesian.

Return ONLY valid JSON with this exact structure:
{{
  "meeting_name": "string",
  "meeting_date": "string",
  "meeting_time": "string",
  "meeting_location": "string",
  "chairperson": "string",
  "prepared_by": "string",
  "next_meeting_date": "string",
  "next_meeting_time": "string",
  "next_meeting_location": "string",
  "next_meeting_chairperson": "string",
  "next_meeting_agenda": "string",
  "items": [
    {{
      "issue": "string - action item atau keputusan dalam Bahasa Indonesia",
      "deadline": "string (dd/mm/yyyy atau TBA atau Note)",
      "status": "string (Open/On Progress/Done/Monitoring/Note)",
      "pic": "string - nama penanggung jawab"
    }}
  ]
}}

Meeting name hint: {meeting_name}

Transcript:
{transcript[:8000]}"""

    try:
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=4000,
        )
        raw = resp.choices[0].message.content.strip()
        # strip markdown code block if present
        raw = re.sub(r'^```(?:json)?\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        return json.loads(raw)
    except json.JSONDecodeError as e:
        return {"error": f"AI returned invalid JSON: {e}", "raw": raw}
    except Exception as e:
        err_str = str(e)
        if 'rate_limit' in err_str.lower() or '429' in err_str:
            retry_match = re.search(r'try again in ([^\.\n]+)', err_str, re.IGNORECASE)
            retry_info = retry_match.group(1).strip() if retry_match else None
            msg = "Kuota AI habis (rate limit)."
            if retry_info:
                msg += f" Coba lagi dalam: {retry_info}."
            else:
                msg += " Coba lagi dalam beberapa menit."
            return {"error": msg, "error_type": "rate_limit"}
        return {"error": err_str}


def slugify(text: str) -> str:
    text = re.sub(r'[^\w\s-]', '', text.lower())
    return re.sub(r'[\s_]+', '-', text).strip('-')


def cm(val): return Inches(val / 2.54)

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _cell_text(cell, text, font_name='Calibri', font_size=9, bold=False,
               italic=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def _set_borders(cell, color='000000'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _apply_borders_table(table, color='000000'):
    for row in table.rows:
        for cell in row.cells:
            _set_borders(cell, color)

def _set_min_row_height(row, height_cm: float):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

def _apply_min_row_height(table, height_cm=0.48):
    for row in table.rows:
        _set_min_row_height(row, height_cm)

STATUS_COLORS = {
    'monitoring': '00FFFF',
    'note': '00FFFF',
    'notes': '00FFFF',
    'noted': '00FFFF',
    'on progress': 'FFFF00',
    'progress': 'FFFF00',
    'open': 'FFFF00',
    'done': '00FF00',
}

def _status_bg(status: str) -> str:
    return STATUS_COLORS.get(status.lower().strip(), 'FFFFFF')

LOGO_PATH = Path(__file__).parent.parent / 'logo_dbc.png'

def _add_header_footer(doc):
    section = doc.sections[0]

    # ── HEADER ────────────────────────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    ht = header.add_table(1, 2, width=cm(18.3))
    ht.autofit = False

    # Set header row height to exactly 1.44cm
    hrow = ht.rows[0]
    trPr = hrow._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(1.44 * 567)))  # 1cm = 567 twips
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)

    logo_cell = hrow.cells[0]
    _set_cell_bg(logo_cell, 'FFFFFF')
    _set_borders(logo_cell, 'C0C0C0')
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Logo height = row height (1.44cm), width proportional
    if LOGO_PATH.exists():
        lp.add_run().add_picture(str(LOGO_PATH), height=cm(1.44))

    title_cell = hrow.cells[1]
    _set_cell_bg(title_cell, 'FFFFFF')
    _set_borders(title_cell, 'C0C0C0')
    _cell_text(title_cell, 'Minute of Meeting', font_name='Arial',
               font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_field(run, field_name):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld)
        instr = OxmlElement('w:instrText')
        instr.text = field_name
        run._r.append(instr)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld2)

    for txt, field in [('Page ', 'PAGE'), (' of ', None), ('', 'NUMPAGES')]:
        r = fp.add_run(txt)
        r.font.name = 'Arial'
        r.font.size = Pt(8)
        if field:
            _add_field(r, field)


def build_docx(data: dict, image_paths: list = None) -> io.BytesIO:
    doc = Document()

    # ── Page: Letter Portrait ─────────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = cm(21.59)
        sec.page_height   = cm(27.94)
        sec.top_margin    = cm(2.54)
        sec.bottom_margin = cm(2.54)
        sec.left_margin   = cm(1.25)
        sec.right_margin  = cm(2.54)

    _add_header_footer(doc)

    # ── 1. MEETING INFORMATION ────────────────────────────────────────────────
    t1 = doc.add_table(rows=10, cols=4)
    t1.style = 'Table Grid'
    t1.autofit = False

    def _info_header(row_idx, label):
        r = t1.rows[row_idx]
        c = r.cells[0].merge(r.cells[3])
        _set_cell_bg(c, '4C4C4C')
        _cell_text(c, label, bold=True, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_borders(c)

    def _info_row(row_idx, label, value, label2=None, value2=None):
        r = t1.rows[row_idx]
        _set_cell_bg(r.cells[0], 'E6E6E6')
        _cell_text(r.cells[0], label)
        _set_borders(r.cells[0])
        if label2:
            _cell_text(r.cells[1], value)
            _set_borders(r.cells[1])
            _set_cell_bg(r.cells[2], 'E6E6E6')
            _cell_text(r.cells[2], label2)
            _set_borders(r.cells[2])
            _cell_text(r.cells[3], value2 or '')
            _set_borders(r.cells[3])
        else:
            c = r.cells[1].merge(r.cells[3])
            _cell_text(c, value)
            _set_borders(c)

    attendees = data.get('attendees', [])
    meeting_dt = data.get('meeting_date', '')
    if data.get('meeting_time'):
        meeting_dt += f" ({data['meeting_time']})"

    _info_header(0, 'Meeting Information')
    _info_row(1, 'Group / Meeting / Project Name', data.get('meeting_name', ''))
    _info_row(2, 'Meeting Date', meeting_dt, 'Meeting Location', data.get('meeting_location', ''))
    _info_row(3, 'Chairperson', data.get('chairperson', ''))
    _info_row(4, 'Attendees', ', '.join(attendees) if attendees else '')
    _info_row(5, 'With Apologies', '')
    _info_row(6, 'Doc Circulation', '', 'Prepared By', data.get('prepared_by', ''))
    _info_row(7, 'Additional Info', '')

    _apply_min_row_height(t1)

    # ── 2. TOPIC & ATTACHMENT ─────────────────────────────────────────────────
    doc.add_paragraph()
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    t2.autofit = False
    r = t2.rows[0]
    _set_cell_bg(r.cells[0], 'E6E6E6')
    _cell_text(r.cells[0], 'Topic')
    _set_borders(r.cells[0])
    _cell_text(r.cells[1], data.get('meeting_name', ''))
    _set_borders(r.cells[1])
    _set_cell_bg(r.cells[2], 'E6E6E6')
    _cell_text(r.cells[2], 'Attachment')
    _set_borders(r.cells[2])
    _cell_text(r.cells[3], '-')
    _set_borders(r.cells[3])

    _apply_min_row_height(t2)

    # ── 3. ISSUES AND ACTIONS ─────────────────────────────────────────────────
    items = data.get('items', [])
    t3 = doc.add_table(rows=1 + len(items), cols=4)
    t3.style = 'Table Grid'
    t3.autofit = False

    # Header row
    rh = t3.rows[0]
    for ci, (label, align) in enumerate([
        ('Issues and Actions', WD_ALIGN_PARAGRAPH.CENTER),
        ('Deadline',           WD_ALIGN_PARAGRAPH.CENTER),
        ('Status',             WD_ALIGN_PARAGRAPH.CENTER),
        ('PIC',                WD_ALIGN_PARAGRAPH.CENTER),
    ]):
        _set_cell_bg(rh.cells[ci], 'E6E6E6')
        _cell_text(rh.cells[ci], label, bold=True, italic=True, align=align)
        _set_borders(rh.cells[ci])

    # Set column widths (approx): Issues 10cm, Deadline 3cm, Status 3cm, PIC 2cm
    for row in t3.rows:
        for ci, w in enumerate([10, 3, 3, 2]):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w * 567))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # Enable header row repeat
    trPr = t3.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

    # Item rows
    for idx, item in enumerate(items):
        ri = t3.rows[1 + idx]
        status = item.get('status', '')
        bg = _status_bg(status)

        _cell_text(ri.cells[0], item.get('issue', ''), align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_borders(ri.cells[0])
        _set_cell_bg(ri.cells[0], 'FFFFFF')

        _cell_text(ri.cells[1], item.get('deadline', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_borders(ri.cells[1])
        _set_cell_bg(ri.cells[1], 'FFFFFF')

        _cell_text(ri.cells[2], status, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_borders(ri.cells[2])
        _set_cell_bg(ri.cells[2], bg)

        _cell_text(ri.cells[3], item.get('pic', ''), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_borders(ri.cells[3])
        _set_cell_bg(ri.cells[3], 'FFFFFF')

    _apply_min_row_height(t3)

    # ── 4. TENTATIVE AGENDA FOR THE NEXT MEETING ──────────────────────────────
    doc.add_paragraph()
    t4 = doc.add_table(rows=5, cols=4)
    t4.style = 'Table Grid'
    t4.autofit = False

    # Header
    hc = t4.rows[0].cells[0].merge(t4.rows[0].cells[3])
    _set_cell_bg(hc, '000000')
    _cell_text(hc, 'Tentative Agenda for the Next Meeting',
               bold=True, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_borders(hc)

    def _agenda_row(row_idx, label, value, label2=None, value2=None):
        r = t4.rows[row_idx]
        _set_cell_bg(r.cells[0], 'D9D9D9')
        _cell_text(r.cells[0], label)
        _set_borders(r.cells[0])
        if label2:
            _cell_text(r.cells[1], value)
            _set_borders(r.cells[1])
            _set_cell_bg(r.cells[2], 'D9D9D9')
            _cell_text(r.cells[2], label2)
            _set_borders(r.cells[2])
            _cell_text(r.cells[3], value2 or '')
            _set_borders(r.cells[3])
        else:
            c = r.cells[1].merge(r.cells[3])
            _cell_text(c, value)
            _set_borders(c)

    def _nm(val): return '-' if not val or val.strip().upper() == 'TBA' else val
    _agenda_row(1, 'Proposed Date',  _nm(data.get('next_meeting_date', '')),
                   'Proposed Location', _nm(data.get('next_meeting_location', '')))
    _agenda_row(2, 'Proposed Time',  _nm(data.get('next_meeting_time', '')),
                   'Chairperson',       _nm(data.get('next_meeting_chairperson', '')))
    _agenda_row(3, 'Proposed Agenda', _nm(data.get('next_meeting_agenda', '')))
    _agenda_row(4, 'Additional Info', '')

    _apply_min_row_height(t4)

    # ── 5. LAMPIRAN GAMBAR ────────────────────────────────────────────────────
    if image_paths:
        doc.add_page_break()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run('Lampiran')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = True

        # Available width: Letter - margins = 21.59 - 1.25 - 2.54 = 17.8cm
        max_w = cm(17.8)

        for idx, img_item in enumerate(image_paths):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_item['path'], width=max_w)
            except Exception:
                continue

            caption_text = img_item.get('caption') or f'Gambar {idx + 1}'
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption_text)
            run_cap.font.name = 'Calibri'
            run_cap.font.size = Pt(9)
            run_cap.italic = True
            run_cap.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── routes ────────────────────────────────────────────────────────────────────

def _compress_images(img_paths: list, quality: int = 70, max_width: int = 1200) -> list:
    """Compress images in-place to JPEG, return new paths."""
    tmp = tempfile.mkdtemp()
    new_paths = []
    for p in img_paths:
        try:
            img = PILImage.open(p)
            if img.width > max_width:
                ratio = max_width / img.width
                img = img.resize((max_width, int(img.height * ratio)), PILImage.LANCZOS)
            img = img.convert("RGB")
            out = Path(tmp) / (Path(p).stem + ".jpg")
            img.save(str(out), "JPEG", quality=quality, optimize=True)
            new_paths.append(str(out))
        except Exception:
            new_paths.append(p)
    return new_paths, tmp


@app.route('/api/upload-image', methods=['POST'])
def api_upload_image():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': 'Empty filename'}), 400
    # Save with unique name
    ext = Path(f.filename).suffix.lower()
    if ext not in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
        return jsonify({'error': 'Format tidak didukung'}), 400
    uid = f"{datetime.now().strftime('%Y%m%d%H%M%S%f')}{ext}"
    save_path = UPLOAD_DIR / uid
    f.save(str(save_path))
    return jsonify({'id': uid, 'name': f.filename})


@app.route('/api/delete-image/<uid>', methods=['DELETE'])
def api_delete_image(uid):
    p = UPLOAD_DIR / uid
    if p.exists():
        p.unlink()
    return jsonify({'ok': True})


@app.route('/api/image/<uid>')
def api_serve_image(uid):
    return send_from_directory(str(UPLOAD_DIR), uid)

def extract_attendees(text: str) -> list:
    """Extract speaker names from transcript without AI."""
    patterns = [
        r'^\[[\d:]+\]\s+([A-Z][a-zA-Z\s]+?)(?:\s*\(.*?\))?\s*:',
        r'^([A-Z][a-zA-Z\s]+?)(?:\s*\(.*?\))?\s*:',
        r'^([A-Z][a-zA-Z\s]+?)\s+\(\d{2}:\d{2}\)',
    ]
    names = set()
    for line in text.splitlines():
        line = line.strip()
        for p in patterns:
            m = re.match(p, line)
            if m:
                name = m.group(1).strip()
                if 2 < len(name) < 40 and not any(w in name.lower() for w in [
                    'http','www','note','action','recording','solution','proposed',
                    'summary','agenda','topic','discussion','meeting','update',
                    'follow','review','next','view','item','issue','status',
                    'deadline','attachment','information','additional','prepared',
                    'chairperson','circulation','apologies','location','date','time',
                ]):
                    names.add(name)
                break
    return sorted(names)


def extract_from_calendar(text: str) -> list:
    """Extract names/emails from pasted Google Calendar event description."""
    names = set()
    emails_with_names = set()
    
    # Match "Name <email>" or "Name (email)" - with or without quotes
    for m in re.finditer(r'"?([A-Z][a-zA-Z\s.]{2,30})"?\s*[<(]([\w.+-]+@[\w.-]+\.\w+)', text):
        names.add(m.group(1).strip())
        emails_with_names.add(m.group(2))
    
    # Match emails → use local part as name (only if email not already paired with name)
    emails = re.findall(r'[\w.+-]+@[\w.-]+\.\w+', text)
    for email in emails:
        if email not in emails_with_names:
            local = email.split('@')[0]
            # Remove trailing numbers (e.g., agust123 → agust)
            local = re.sub(r'\d+$', '', local)
            # Convert john.doe or john_doe → John Doe
            name = re.sub(r'[._-]', ' ', local).title()
            names.add(name)
    
    return sorted(names)


@app.route('/api/extract-attendees', methods=['POST'])
def api_extract_attendees():
    body = request.get_json()
    transcript = body.get('transcript', '')
    calendar_text = body.get('calendar_text', '')

    names = set()
    if transcript:
        names.update(extract_attendees(transcript))
    if calendar_text:
        names.update(extract_from_calendar(calendar_text))

    return jsonify({'attendees': sorted(names)})


@app.route('/api/generate', methods=['POST'])
def api_generate():
    body = request.get_json()
    transcript = body.get('transcript', '')
    meeting_name = body.get('meeting_name', 'Meeting')
    calendar_text = body.get('calendar_text', '')

    result = generate_mom_with_ai(transcript, meeting_name)
    if 'error' in result:
        return jsonify(result), 500

    # Default location
    if not result.get('meeting_location'):
        result['meeting_location'] = 'Online'
    ai_attendees = set(result.get('attendees') or [])
    ai_attendees.update(extract_attendees(transcript))
    if calendar_text:
        ai_attendees.update(extract_from_calendar(calendar_text))
    result['attendees'] = sorted(ai_attendees)

    date_str = extract_date_from_transcript(transcript)
    result['_date_for_filename'] = date_str

    return jsonify(result)


@app.route('/api/export', methods=['POST'])
def api_export():
    data = request.get_json()

    date_str = data.get('_date_for_filename') or data.get('_manual_date', '')
    if not date_str:
        date_str = datetime.now().strftime('%Y%m%d')

    name_slug = slugify(data.get('meeting_name', 'meeting'))[:40]
    filename = f"{date_str}_{name_slug}.docx"

    # Resolve uploaded image paths
    image_ids = data.get('image_ids', [])
    image_paths = []
    for item in image_ids:
        p = UPLOAD_DIR / item['id']
        if p.exists():
            image_paths.append({'path': str(p), 'caption': item.get('caption', '')})

    buf = build_docx(data, image_paths)

    # Compress if >15MB
    if buf.getbuffer().nbytes > MAX_DOCX_SIZE and image_paths:
        tmp_dirs = []
        for quality, max_width in [(70, 1200), (50, 1000), (35, 800), (20, 600)]:
            raw_paths = [i['path'] for i in image_paths]
            compressed, tmp = _compress_images(raw_paths, quality, max_width)
            tmp_dirs.append(tmp)
            compressed_image_paths = [
                {'path': cp, 'caption': image_paths[i]['caption']}
                for i, cp in enumerate(compressed)
            ]
            buf = build_docx(data, compressed_image_paths)
            if buf.getbuffer().nbytes <= MAX_DOCX_SIZE:
                break
        for t in tmp_dirs:
            shutil.rmtree(t, ignore_errors=True)

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'groq_configured': bool(GROQ_API_KEY)})


if __name__ == '__main__':
    app.run(port=5000, debug=False)
